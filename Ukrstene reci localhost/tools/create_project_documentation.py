from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "Dokumentacija_Ukrstene_reci.docx"
LOGO = ROOT / "public" / "ukrstene-logo.png"
SCREENSHOT = ROOT / "public" / "ukrstene-gameplay-preview.jpg"
GITHUB_URL = "https://github.com/lazarleka/Wordsearch-Game"


def set_run(run, size=None, bold=False, color=None):
    run.font.name = "Calibri"
    if size:
        run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    run = p.add_run(text)
    set_run(run, 16 if level == 1 else 13, True, (46, 116, 181) if level == 1 else (31, 77, 120))
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    set_run(run, 11)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    run = p.add_run(text)
    set_run(run, 11)
    return p


def add_body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.1
    run = p.add_run(text)
    set_run(run, 11)
    return p


def build_doc():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if LOGO.exists():
        title.add_run().add_picture(str(LOGO), width=Inches(1.0))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Dokumentacija projekta: Ukrstene rijeci")
    set_run(run, 22, True, (46, 116, 181))

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Web igra osmosmjerke sa AI generisanjem rijeci, solo i versus modom")
    set_run(run, 11, False, (85, 85, 85))

    add_heading(doc, "1. Opis igre", 1)
    add_body(
        doc,
        "Ukrstene rijeci je web igra zasnovana na osmosmjerci. Igrac bira temu i tezinu, "
        "a zatim na tabli trazi skrivene rijeci. Rijeci se mogu nalaziti horizontalno, "
        "vertikalno i dijagonalno, u oba smjera. Cilj je pronaci sto vise rijeci za sto krace vrijeme.",
    )
    add_bullet(doc, "Solo mod: igrac igra sam, skuplja bodove i rezultat ulazi u istoriju i rang listu.")
    add_bullet(doc, "Multiplayer mod: dva igraca igraju lokalno na istom uredjaju i smjenjuju poteze.")
    add_bullet(doc, "Versus mod: registrovani korisnik salje izazov prijatelju i oba igraca igraju istu tablu.")
    add_bullet(doc, "Rang lista i istorija meceva prikazuju napredak i rezultate korisnika.")
    add_bullet(doc, "Tokom igre postoje zvucni efekti za pogodak, promasaj i zavrsetak partije.")

    add_heading(doc, "2. Inovacije u odnosu na original", 1)
    add_body(
        doc,
        "Originalna igra osmosmjerke uglavnom podrazumijeva staticku tablu i unaprijed pripremljene rijeci. "
        "Ova verzija prosiruje ideju kroz personalizovano generisanje sadrzaja, naloge, takmicenje i moderniji tok igre.",
    )
    add_bullet(doc, "AI generisanje rijeci prema temi koju korisnik izabere ili unese.")
    add_bullet(doc, "Korisnicki nalozi, login/registracija, prijatelji i zahtjevi za prijateljstvo.")
    add_bullet(doc, "Versus izazovi izmedju prijatelja sa istim setom rijeci za oba igraca.")
    add_bullet(doc, "Solo partije se zavrsavaju nakon 5 minuta ili ranije ako igrac pronadje sve rijeci.")
    add_bullet(doc, "Power-up opcije: prikaz prvog slova i pomocna putanja, uz kaznu u bodovima.")
    add_bullet(doc, "Rang lista obuhvata rezultate iz solo i versus partija.")
    add_bullet(doc, "Pozadinska muzika sa mute opcijom i odvojeni zvucni efekti za uspjeh/promasaj.")
    add_bullet(doc, "Automatski oporavak ako AI vrati nevalidne rijeci, da se korisnicki nalog ne blokira.")

    add_heading(doc, "3. Koristeni AI alati", 1)
    add_body(
        doc,
        "U izradi projekta koristeni su AI alati za generisanje sadrzaja, pomoc pri programiranju i izradu/promjenu vizuelnih i zvucnih materijala.",
    )
    add_number(doc, "Google Gemini API - koristi se u aplikaciji za generisanje tematskih rijeci za osmosmjerku.")
    add_number(doc, "OpenAI/Codex - koristen kao AI asistent za pomoc pri razvoju, ispravljanju gresaka, dokumentovanju i provjeri projekta.")
    add_number(doc, "AI alati za multimediju - koristeni su za izradu ili obradu promotivnih/grafickih materijala i zvucnog identiteta igre.")
    add_body(
        doc,
        "AI generisane rijeci dodatno se validiraju: aplikacija filtrira rijeci koje su krace od 4 ili duze od 12 slova i pokusava ponovo ako set nije pogodan za tablu.",
    )

    add_heading(doc, "4. Screenshotovi igre", 1)
    add_body(doc, "U nastavku je prikazan primjer gameplay ekrana igre.")
    if SCREENSHOT.exists():
        img_p = doc.add_paragraph()
        img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        img_p.add_run().add_picture(str(SCREENSHOT), width=Inches(5.8))
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap.add_run("Slika 1. Primjer table i rijeci u igri Ukrstene rijeci.")
        set_run(run, 9, False, (85, 85, 85))
    else:
        add_body(doc, "Screenshot nije pronadjen u public folderu. Dodati sliku gameplay ekrana prije predaje.")

    add_heading(doc, "5. Link do repozitorijuma", 1)
    add_body(doc, f"GitHub repozitorijum projekta: {GITHUB_URL}")

    add_heading(doc, "6. Zakljucak", 1)
    add_body(
        doc,
        "Projekat Ukrstene rijeci predstavlja prosirenu verziju klasicne osmosmjerke. "
        "Pored osnovnog trazenja rijeci, igra ukljucuje AI generisanje, vise modova igre, takmicenje, rang listu, "
        "zvucne efekte, pozadinsku muziku i sistem prijatelja. Time projekat zadovoljava zahtjeve za inovativnom, "
        "multimedijalnom i AI podrzanom verzijom originalne igre.",
    )

    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    print(build_doc())
