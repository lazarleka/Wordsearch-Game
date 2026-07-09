from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"C:\Users\User\OneDrive\Desktop\grafika\Ukrstene reci localhost")
OUT = Path(r"C:\Users\User\OneDrive\Desktop\grafika\Dokumentacija_unapredjena.docx")
LOGO = ROOT / "public" / "ukrstene-logo.png"
GAMEPLAY = ROOT / "public" / "ukrstene-gameplay-preview.jpg"


BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
MUTED = "6B7280"
LIGHT = "E8EEF5"
LIGHTER = "F4F6F9"
BORDER = "C9D3DF"
TEAL = "00A98F"
GOLD = "B7791F"


def set_font(run, name="Calibri", size=11, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_p_spacing(paragraph, before=0, after=6, line=1.25):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color=BORDER, size="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    mar = tc_pr.first_child_found_in("w:tcMar")
    if mar is None:
        mar = OxmlElement("w:tcMar")
        tc_pr.append(mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = mar.find(qn("w:" + m))
        if node is None:
            node = OxmlElement("w:" + m)
            mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths)))
    ind = tbl_pr.find(qn("w:tblInd"))
    if ind is None:
        ind = OxmlElement("w:tblInd")
        tbl_pr.append(ind)
    ind.set(qn("w:type"), "dxa")
    ind.set(qn("w:w"), "120")

    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = tr_pr.find(qn("w:cantSplit"))
        if cant_split is None:
            tr_pr.append(OxmlElement("w:cantSplit"))
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(widths[idx]))
            set_cell_margins(cell)
            set_cell_border(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def keep_with_next(paragraph):
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.keep_together = True


def add_para(doc, text="", style=None, bold_prefix=None):
    p = doc.add_paragraph(style=style)
    set_p_spacing(p)
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_font(r, bold=True, color=INK)
        r2 = p.add_run(text[len(bold_prefix):])
        set_font(r2)
    else:
        r = p.add_run(text)
        set_font(r)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    set_p_spacing(p, after=4)
    if ": " in text:
        head, rest = text.split(": ", 1)
        r = p.add_run(head + ": ")
        set_font(r, bold=True, color=INK)
        r = p.add_run(rest)
        set_font(r)
    else:
        set_font(p.add_run(text))
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    set_p_spacing(p, after=4)
    set_font(p.add_run(text))
    return p


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    keep_with_next(p)
    if level == 1:
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(10)
        size, color = 16, BLUE
    elif level == 2:
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(7)
        size, color = 13, BLUE
    else:
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(5)
        size, color = 12, DARK_BLUE
    for run in p.runs:
        set_font(run, size=size, color=color, bold=True)
    return p


def add_label_table(doc, rows):
    table = doc.add_table(rows=len(rows), cols=2)
    set_table_width(table, [2700, 6660])
    for i, (label, value) in enumerate(rows):
        left, right = table.rows[i].cells
        shade_cell(left, LIGHT)
        for cell in (left, right):
            for p in cell.paragraphs:
                set_p_spacing(p, after=0)
        left.paragraphs[0].add_run(label)
        set_font(left.paragraphs[0].runs[0], bold=True, color=INK)
        right.paragraphs[0].add_run(value)
        set_font(right.paragraphs[0].runs[0])
    doc.add_paragraph()
    return table


def add_matrix_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    set_table_width(table, widths)
    for idx, h in enumerate(headers):
        cell = table.rows[0].cells[idx]
        shade_cell(cell, LIGHT)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(h)
        set_font(p.runs[0], size=10.5, bold=True, color=INK)
    for r_idx, row in enumerate(rows, 1):
        for c_idx, value in enumerate(row):
            cell = table.rows[r_idx].cells[c_idx]
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.1
            p.add_run(value)
            set_font(p.runs[0], size=9.8)
            if c_idx == 0:
                p.runs[0].bold = True
                p.runs[0].font.color.rgb = RGBColor.from_string(INK)
    doc.add_paragraph()
    return table


def add_callout(doc, title, text, fill=LIGHTER, color=INK):
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, [9360])
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    p = cell.paragraphs[0]
    set_p_spacing(p, after=2)
    r = p.add_run(title)
    set_font(r, bold=True, color=color)
    p2 = cell.add_paragraph()
    set_p_spacing(p2, after=0)
    r = p2.add_run(text)
    set_font(r, size=10.5)
    doc.add_paragraph()


def setup_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name in ("List Bullet", "List Number"):
        st = doc.styles[style_name]
        st.font.name = "Calibri"
        st.font.size = Pt(11)
        st.paragraph_format.space_after = Pt(4)
        st.paragraph_format.line_spacing = 1.25


def build():
    doc = Document()
    setup_styles(doc)

    # Cover/title block.
    if LOGO.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(LOGO), width=Inches(0.85))
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(4)
    title.paragraph_format.space_after = Pt(4)
    r = title.add_run("Dokumentacija projekta: Ukrštene riječi")
    set_font(r, size=24, color=INK, bold=True)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(12)
    r = subtitle.add_run("Web i Android igra osmosmjerke sa AI generisanjem riječi, nalozima i takmičarskim modovima")
    set_font(r, size=12, color=MUTED)

    add_label_table(doc, [
        ("Tip projekta", "Interaktivna igra / web aplikacija / Android aplikacija"),
        ("Repozitorijum", "https://github.com/lazarleka/Wordsearch-Game"),
        ("Frontend", "React 18, Vite, Lucide React, responsive CSS"),
        ("Backend", "Java Spring Boot, REST API, WebSocket komunikacija"),
        ("Baza podataka", "MySQL, šema ukrstene_reci"),
        ("AI komponenta", "Google Gemini API za generisanje tematskih riječi"),
        ("Mobilna aplikacija", "Capacitor Android, lokalne notifikacije za izazove"),
    ])

    add_callout(
        doc,
        "Kratak opis",
        "Ukrštene riječi su modernizovana osmosmjerka u kojoj korisnik bira temu i težinu, pronalazi skrivene riječi na tabli i upoređuje rezultate kroz solo, lokalni multiplayer, versus i race mod. Projekat spaja klasičnu logičku igru, AI generisanje sadržaja, korisničke naloge, prijatelje, rang listu i mobilno iskustvo.",
        fill="F6FAFF",
        color=BLUE,
    )

    if GAMEPLAY.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(GAMEPLAY), width=Inches(5.6))
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_after = Pt(10)
        r = cap.add_run("Slika 1: Primjer izgleda table i liste riječi u igri.")
        set_font(r, size=9, color=MUTED, italic=True)

    add_heading(doc, "1. Svrha i cilj projekta")
    add_para(doc, "Cilj projekta je da klasičnu osmosmjerku pretvori u savremenu, personalizovanu i takmičarsku aplikaciju. Umjesto unaprijed zadatih tabli, aplikacija može koristiti riječi iz baze ili AI generisane riječi prema temi koju korisnik izabere ili sam unese.")
    add_para(doc, "Projekat je razvijen tako da pokrije kompletan tok rada: registraciju korisnika, izbor igre, generisanje table, praćenje rezultata, online izazove, rang listu, admin upravljanje temama i mobilnu upotrebu.")

    add_heading(doc, "2. Osnovna pravila igre")
    add_bullet(doc, "Izbor teme: korisnik bira postojeću temu iz baze ili unosi sopstvenu temu za AI generisanje.")
    add_bullet(doc, "Izbor težine: težina određuje veličinu matrice, broj riječi i maksimalnu dužinu riječi.")
    add_bullet(doc, "Traženje riječi: riječi se mogu nalaziti horizontalno, vertikalno i dijagonalno, u oba smjera.")
    add_bullet(doc, "Bodovanje: pronađena riječ donosi bodove, dok power-up opcije smanjuju ukupan rezultat za unaprijed definisanu kaznu.")
    add_bullet(doc, "Završetak partije: igra se završava pronalaskom svih riječi, istekom vremena ili ručnim završetkom/napuštanjem meča.")

    add_heading(doc, "3. Modovi igre")
    add_matrix_table(
        doc,
        ["Mod", "Opis", "Rezultat"],
        [
            ("Solo", "Jedan korisnik igra protiv vremena. Partija traje do 5 minuta ili kraće ako pronađe sve riječi.", "Bodovi se čuvaju u istoriji i utiču na rang listu."),
            ("Multiplayer", "Dva igrača igraju lokalno na istom uređaju i smjenjuju poteze.", "Rezultat se vodi po igraču."),
            ("Versus", "Prijatelji igraju istu tablu kroz online izazov.", "Pobjednik se određuje po učinku u meču."),
            ("Race", "Dva igrača istovremeno osvajaju riječi na istoj tabli.", "Riječ pripada igraču koji je prvi pronađe."),
        ],
        [1500, 5360, 2500],
    )

    add_heading(doc, "4. Vrste table")
    add_bullet(doc, "Obična tabla: svaka riječ je postavljena pravolinijski u jednom od osam smjerova.")
    add_bullet(doc, "Zmijica: riječ može skretati od slova do slova kroz susjedna polja. Putanja je jedinstvena i nakon pronalaska dobija redne oznake slova.")
    add_bullet(doc, "Validacija table: algoritam pokušava više generisanja i odbacuje table na kojima riječi nije moguće korektno postaviti.")

    add_heading(doc, "5. Inovacije u odnosu na klasičnu osmosmjerku")
    add_bullet(doc, "AI generisanje sadržaja: korisnik može unijeti slobodnu temu, a aplikacija generiše odgovarajuće riječi.")
    add_bullet(doc, "Dinamička validacija riječi: filtriraju se riječi kraće od 4 ili duže od dozvoljene granice, riječi sa nepodržanim slovima i neupotrebljivi AI odgovori.")
    add_bullet(doc, "Online izazovi: prijatelji mogu slati i prihvatati izazove sa istom temom, težinom i setom riječi.")
    add_bullet(doc, "Race mehanika: riječ se može osvojiti samo jednom, što dodaje brzinu i taktički pritisak.")
    add_bullet(doc, "Power-up sistem: prvo slovo i pomoćna putanja olakšavaju igru, ali smanjuju bodove.")
    add_bullet(doc, "Mobilne notifikacije: Android aplikacija obavještava korisnika o novim izazovima.")
    add_bullet(doc, "Admin panel: administrator može dodavati, mijenjati i brisati teme i riječi iz baze.")

    add_heading(doc, "6. Korisničke uloge i funkcionalnosti")
    add_matrix_table(
        doc,
        ["Uloga", "Mogućnosti"],
        [
            ("Neregistrovani korisnik", "Može pristupiti ekranu za prijavu i registraciju."),
            ("Registrovani korisnik", "Igra solo i online modove, šalje zahtjeve za prijateljstvo, prihvata izazove, prati istoriju i rang listu."),
            ("Administrator", "Upravlja temama, riječima i predlozima tema kroz admin panel."),
        ],
        [2400, 6960],
    )

    doc.add_page_break()
    add_heading(doc, "7. Arhitektura sistema")
    add_para(doc, "Aplikacija je podijeljena na klijentski dio, serverski dio i bazu podataka. Frontend komunicira sa backendom preko REST API-ja, dok se promjene u online mečevima i izazovima dopunjuju WebSocket kanalima i periodičnim osvježavanjem.")
    add_matrix_table(
        doc,
        ["Sloj", "Tehnologije", "Odgovornost"],
        [
            ("Frontend", "React, Vite, CSS", "Ekrani za prijavu, izbor igre, tablu, rang listu, prijatelje, izazove i admin panel."),
            ("Mobilni omotač", "Capacitor Android", "Pakovanje web aplikacije kao Android aplikacije i rad sa lokalnim notifikacijama."),
            ("Backend", "Spring Boot, REST, WebSocket", "Autentifikacija, mečevi, izazovi, prijatelji, rang lista, admin operacije i podaci za igru."),
            ("Baza", "MySQL", "Korisnici, teme, riječi, prijateljstva, izazovi, mečevi, solo rezultati i predlozi tema."),
            ("AI servis", "Google Gemini API", "Generisanje riječi za korisnički unesene teme kada se ne koristi statična baza."),
        ],
        [1700, 2500, 5160],
    )

    add_heading(doc, "8. Backend API i tok podataka")
    add_bullet(doc, "Autentifikacija: registracija i prijava korisnika, uz čuvanje hash lozinke u bazi.")
    add_bullet(doc, "Teme i riječi: dohvat tema, dohvat riječi po temi i predlaganje novih tema.")
    add_bullet(doc, "Prijatelji: slanje zahtjeva, prihvatanje zahtjeva i prikaz liste prijatelja.")
    add_bullet(doc, "Izazovi i mečevi: kreiranje izazova, prihvatanje/odbijanje, aktivni meč, napredak, završetak i predaja.")
    add_bullet(doc, "Rang lista i istorija: pregled solo rezultata, online mečeva i ukupnih pobjeda/poraza.")
    add_bullet(doc, "Admin API: dodavanje i uređivanje tema i riječi, pregled predloga i moderacija sadržaja.")

    doc.add_page_break()
    add_heading(doc, "9. Model baze podataka")
    add_matrix_table(
        doc,
        ["Tabela", "Svrha"],
        [
            ("korisnik", "Nalozi, uloge, avatar boja i ukupan učinak korisnika."),
            ("tema", "Tematske kategorije koje mogu biti statične ili dodate kroz admin tok."),
            ("tema_rijec", "Pojedinačne riječi povezane sa temama i težinom."),
            ("prijateljstvo", "Status odnosa između korisnika."),
            ("izazov", "Poziv na online meč, tema, težina, tip table i status odgovora."),
            ("mec / mec_igrac", "Stanje online meča, igrači, pronađene riječi, bodovi i pobjednik."),
            ("solo_rezultat", "Rezultati solo partija za istoriju i rangiranje."),
            ("tema_predlog", "Korisnički predlozi novih tema koje administrator može odobriti ili odbiti."),
        ],
        [2300, 7060],
    )

    add_heading(doc, "10. AI alati i njihova uloga")
    add_number(doc, "Google Gemini API se koristi u samoj aplikaciji za generisanje tematskih riječi na osnovu izabrane ili unesene teme.")
    add_number(doc, "OpenAI/Codex je korišten kao razvojni asistent za programiranje, refaktorisanje, ispravljanje grešaka i pripremu dokumentacije.")
    add_number(doc, "AI alati za multimediju korišteni su za pripremu i doradu promotivnih grafika, animacija i zvučnog identiteta.")
    add_callout(
        doc,
        "Kontrola kvaliteta AI sadržaja",
        "AI odgovor se ne koristi direktno bez provjere. Aplikacija normalizuje riječi, uklanja nepodržane karaktere, provjerava dužinu, odbacuje nedovoljan broj validnih riječi i po potrebi ponavlja generisanje.",
        fill="FFF8E6",
        color=GOLD,
    )

    add_heading(doc, "11. Mobilna aplikacija")
    add_para(doc, "Projekat se može pokrenuti kao Android aplikacija preko Capacitor-a. Mobilni build koristi posebne environment varijable za backend URL, API URL, auth URL i WebSocket URL.")
    add_bullet(doc, "Telefon i računar moraju biti na istoj Wi-Fi mreži kada se koristi lokalni backend.")
    add_bullet(doc, "Backend mora slušati na mrežnom interfejsu dostupnom telefonu, a Windows Firewall mora dozvoliti port 8082.")
    add_bullet(doc, "Komanda npm run mobile:apk pravi debug APK u android/app/build/outputs/apk/debug/app-debug.apk.")
    add_bullet(doc, "Lokalne notifikacije provjeravaju nove izazove i vode korisnika na karticu Izazovi.")

    add_heading(doc, "12. Lokalno pokretanje")
    add_matrix_table(
        doc,
        ["Komanda / servis", "Namjena"],
        [
            ("npm run dev:full", "Pokreće lokalnu pripremu baze, backend i frontend preko start-local.ps1."),
            ("npm run dev:backend", "Pokreće samo Spring Boot backend."),
            ("npm run dev:frontend", "Pokreće Vite frontend za razvoj."),
            ("npm run build", "Pravi produkcijski frontend build."),
            ("npm run mobile:sync", "Pravi mobilni build i sinhronizuje ga sa Android projektom."),
            ("npm run mobile:apk", "Generiše Android debug APK."),
        ],
        [2800, 6560],
    )
    add_para(doc, "Podrazumijevani lokalni servisi: frontend http://localhost:4713, backend http://localhost:8082, MySQL localhost:3306, baza ukrstene_reci.")

    add_heading(doc, "13. Testiranje i provjera")
    add_bullet(doc, "Funkcionalno testiranje: registracija, prijava, izbor teme, start igre, pronalazak riječi i završetak partije.")
    add_bullet(doc, "Testiranje modova: solo, lokalni multiplayer, versus, race, obična tabla i zmijica.")
    add_bullet(doc, "Testiranje AI toka: validna tema, nedovoljno riječi, ponovljeno generisanje i fallback poruke.")
    add_bullet(doc, "Testiranje online toka: zahtjev za prijateljstvo, izazov, prihvatanje, napredak meča i završetak.")
    add_bullet(doc, "Responsivni test: desktop, mobilni browser i Android aplikacija, naročito raspored informacija tokom igre.")
    add_bullet(doc, "Baza: provjera kreiranja tabela, stranih ključeva, solo rezultata, mečeva i admin promjena.")

    add_heading(doc, "14. Moguća buduća unapređenja")
    add_bullet(doc, "Dodati javne dnevne izazove sa istom temom za sve korisnike.")
    add_bullet(doc, "Uvesti detaljniju statistiku po temi, težini i prosječnom vremenu pronalaska riječi.")
    add_bullet(doc, "Dodati globalne achievement-e i sezonske rang liste.")
    add_bullet(doc, "Omogućiti offline solo mod sa lokalno sačuvanim rezultatima koji se kasnije sinhronizuju.")
    add_bullet(doc, "Dodati administratorski uvoz riječi iz CSV/XLSX fajla.")
    add_bullet(doc, "Unaprijediti pristupačnost većim kontrastom, opcijama za veličinu slova i režimom bez animacija.")

    add_heading(doc, "15. Zaključak")
    add_para(doc, "Ukrštene riječi predstavljaju proširenu verziju klasične osmosmjerke u kojoj se tradicionalna logička igra kombinuje sa modernim web tehnologijama, AI generisanjem sadržaja, takmičarskim modovima i mobilnim iskustvom. Projekat demonstrira rad sa frontend aplikacijom, backend servisima, bazom podataka, real-time komunikacijom, korisničkim nalozima i administracijom sadržaja.")
    add_para(doc, "Najveća vrijednost projekta je u tome što isti osnovni koncept igre podržava više scenarija: brzo solo igranje, lokalno takmičenje, online izazove sa prijateljima, race mod i personalizovane teme. Time dokumentovani sistem ima jasnu funkcionalnu, tehničku i prezentacionu cjelinu.")

    section = doc.sections[0]
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(footer.add_run("Ukrštene riječi - projektna dokumentacija"), size=9, color=MUTED)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
